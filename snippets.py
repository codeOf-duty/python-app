import docx
import hashlib
import pandas as pd
import os
import subprocess
import platform
import sys  
import shutil
import random as rd
import file_path as fp
from sympound import sympound
from pyxdameraulevenshtein import damerau_levenshtein_distance


distancefun = damerau_levenshtein_distance
ssc = sympound(distancefun=distancefun, maxDictionaryEditDistance=3)
ssc.load_dictionary(fp.SymspellModel, term_index=0, count_index=1)

def remove_pun(line):
  punc = "!#$%&'()*+,-./:‘’’’;<=–>?@[\]^_`{|}~"  
  for ele in line:
      if ele in punc:
          line = line.replace(ele, "\n")
  return line

def convertDoc2Docx(PATH, dest="123"):
    if dest == "123":
        dest = PATH
    for doc in os.listdir(PATH):
        if doc.endswith('.doc'):
            subprocess.call(['soffice', '--headless', '--convert-to', 'docx', '--outdir', dest, PATH+doc])
            os.remove(PATH+doc)

def DocxToList(PATH,opt=True):
  doc = docx.Document(PATH)
  data = []
  for para in doc.paragraphs:
      lst = remove_pun(para.text)
      data.extend(list(map(str, lst.split())))
  newData = []
  for p in data:
    newData.append(p.replace('\u200c',''))
  if opt:
    newData = set(newData)
  return list(newData)



def DocxToData(PATH):
  doc = docx.Document(PATH)
  data = []
  for para in doc.paragraphs:
      data.append(para.text)
  return "\n".join(data)



def ListToDict(data):
  dict_ = {}
  for i in data:
    if i not in dict_:
      id_ = hashlib.md5(i.encode('utf-8')).hexdigest()[:12]
      dict_[id_] = i
  return dict_



def remove_punToList(words):
  pun = "!#$%&'()*+,-./:‘’’;<=>?@[\]^–_`{|}~"
  res = ""
  temp = []
  for i in words:
    if i not in pun:
      res += i
    else:
      temp.append(res)
      res =""
  if len(temp) == 0:
    temp.append(res)
  return temp


def xlxsToList(PATH):
  df = pd.read_excel(PATH) 
  res = df.to_numpy()
  temp = list()
  rej = 0
  wrd = 0
  for i in res:
    for j in i:
      for k in str(j).split():
        wrd += 1
        if k not in temp :
          word = remove_punToList(k)
          for l in word:
            if l not in temp:
              temp.append(l)
              wrd += 1
            else:
              wrd += 1
              rej += 1
        else:
            wrd += 1
            rej += 1
  return temp

def removeLexWord(word_):
    import string
    for c in string.punctuation:
        word_ = word_.replace(c,"")
    for i in ["೦","೧","೨","೩","೪","೫","೬","೭","೮","೯"]:
        word_ = word_.replace(i,"")
    return word_

def suggestionReturner(word__):
    return ssc.lookup_compound(input_string=word__, edit_distance_max=3)

def tagAdder(data, idset):
  for x,y in idset.items():
    if y in data and y not in "!#$%&'()*+,-./:‘’’’;<=–>?@[\]^_`{|}~" and y not in '"':
      temp = removeLexWord(y)
      temp = str(suggestionReturner(temp))
      temp = temp.split(":")[0]
      temp = removeLexWord(temp)
      data = data.replace(y,'<span  class="{}" id="underline" rel="popover" data-toggle="popover" data-trigger="focus" data-html="true"  data-placement="bottom" data-content="<p class=closePop>x</p><p id={} class=SA >{}</p><hr/><p id={} class=A2D  >Add to Dictionary</p><hr/><!--<p id={} {} class=SW  >Suggestion</p><hr/>--><p id={} class=RW  >Replace All</p><hr/><p id={} class=IW  >Ignore</p>"  >{}</span>'.format(x,x,temp,x,x,'data-toggle=collapse href=#suggestionOne',x,x,y))
  #print(data)
  return data
